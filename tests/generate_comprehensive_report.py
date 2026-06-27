import os
import json
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

with open("device_db.json") as f:
    db = json.load(f)

doc = Document()
doc.add_heading('Comprehensive Area Models Document', 0)

# =======================
# 1. MOSFETS
# =======================
doc.add_heading('1. MOSFET Devices', level=1)
doc.add_paragraph(
    "MOSFET area is modeled using a bounding-box approach rather than a flat linear sum. "
    "The bounding box Height depends linearly on Width (W), and the bounding box Width depends on Finger Length (L) and Number of Fingers (nf). "
    "Area is given by the cross-product:\n\n"
    "    Area = (ah*W + bh) * (aw*L*nf + bw*nf + cw)"
)

mosfets = list(db["mosfets"].keys())
for dev in mosfets:
    data = db["mosfets"][dev]
    pts = data.get("sweep", [])
    if len(pts) < 3: continue

    doc.add_heading(f'Device: {dev}', level=2)
    m = data["model"]
    doc.add_paragraph(
        f"Parameters: ah={m['ah']:.4f}, bh={m['bh']:.4f}, aw={m['aw']:.4f}, bw={m['bw']:.4f}, cw={m['cw']:.4f}\n"
        f"Height Model: h_um = {m['ah']:.4f}*W + {m['bh']:.4f}\n"
        f"Width Model: w_um = {m['aw']:.4f}*(L*nf) + {m['bw']:.4f}*nf + {m['cw']:.4f}"
    )

    # Calculate old model parameters to show improvement
    X_old = np.array([[p["w"]*p["nf"], p["l"], 1.] for p in pts])
    y_area = np.array([p["area_um2"] for p in pts])
    try:
        c_old, *_ = np.linalg.lstsq(X_old, y_area, rcond=None)
    except:
        c_old = [0,0,0]

    actual_areas, old_preds, new_preds, labels = [], [], [], []
    
    # Table for document
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Dimensions', 'Actual (µm²)', 'Old Model Err', 'New Model Err'

    for p in pts:
        actual = p["area_um2"]
        pred_old = c_old[0]*(p["w"]*p["nf"]) + c_old[1]*p["l"] + c_old[2]
        pred_new = (m['ah']*p['w'] + m['bh']) * (m['aw']*(p['l']*p['nf']) + m['bw']*p['nf'] + m['cw'])
        
        actual_areas.append(actual)
        old_preds.append(pred_old)
        new_preds.append(pred_new)
        labels.append(f"W={p['w']}\nL={p['l']}\nnf={p['nf']}")
        
        err_old = abs(pred_old - actual) / actual * 100
        err_new = abs(pred_new - actual) / actual * 100
        
        row = table.add_row().cells
        row[0].text = f"W={p['w']} L={p['l']} nf={p['nf']}"
        row[1].text = f"{actual:.2f}"
        row[2].text = f"{pred_old:.2f} ({err_old:.1f}%)"
        row[3].text = f"{pred_new:.2f} ({err_new:.1f}%)"

    # Plot
    plt.figure(figsize=(6, 5))
    min_val = min(min(actual_areas), min(old_preds), min(new_preds)) * 0.9
    max_val = max(max(actual_areas), max(old_preds), max(new_preds)) * 1.1
    plt.plot([min_val, max_val], [min_val, max_val], 'k--', alpha=0.5, label='Perfect Fit')
    plt.scatter(actual_areas, old_preds, color='red', marker='x', s=60, label='Old Linear')
    plt.scatter(actual_areas, new_preds, color='blue', marker='o', s=40, alpha=0.7, label='New Box Model')
    plt.title(f"{dev} Parity Plot")
    plt.xlabel("Actual Magic Area (µm²)")
    plt.ylabel("Predicted Area (µm²)")
    plt.legend()
    plt.grid(True, linestyle=':', alpha=0.6)
    plt.tight_layout()
    img_name = f"plot_{dev}.png"
    plt.savefig(img_name, dpi=120)
    plt.close()

    doc.add_picture(img_name, width=Inches(5.0))
    doc.add_paragraph("")  # spacing

# =======================
# 2. POLY RESISTORS
# =======================
doc.add_heading('2. Poly Resistors', level=1)
doc.add_paragraph(
    "Poly resistor areas are modeled using a linear relationship with length (L) for a given fixed width (W). "
    "Each specific width has its own slope and intercept:\n\n"
    "    Area = slope * L + intercept"
)

for dev, data in db["poly_resistors"].items():
    models = data.get("model", [])
    if not models: continue
    doc.add_heading(f'Device: {dev}', level=2)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'Fixed Width (W)', 'Slope', 'Intercept'
    
    for m in models:
        row = table.add_row().cells
        row[0].text = str(m["w"])
        row[1].text = f"{m['slope']:.4f}"
        row[2].text = f"{m['intercept']:.4f}"
    
    doc.add_paragraph("")

# =======================
# 3. CAPACITORS
# =======================
doc.add_heading('3. Capacitors', level=1)
doc.add_heading('MIM Capacitors', level=2)
doc.add_paragraph(
    "MIM capacitors are modeled as a rectangle with a uniform border overhead (b) around the drawn dimensions:\n\n"
    "    Area = (W + 2b) * (L + 2b)"
)
for dev, data in db["mim_caps"].items():
    m = data.get("model")
    if m:
        doc.add_paragraph(f"- {dev}: border overhead (b) = {m['border_um']:.4f} µm")

doc.add_heading('Varactor Capacitors', level=2)
doc.add_paragraph("Varactors fall back to nearest-neighbor lookup because their scaling is highly non-linear depending on exact array layout.")
for dev, data in db["var_caps"].items():
    pts = data.get("sweep", [])
    doc.add_paragraph(f"- {dev}: {len(pts)} points in database for nearest-neighbor.")

# =======================
# 4. BJTs
# =======================
doc.add_heading('4. Bipolar Junction Transistors (BJTs)', level=1)
doc.add_paragraph("BJTs are fixed-size layout macros. Their areas are static.")
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Device', 'Width (µm)', 'Height (µm)', 'Area (µm²)'

for dev, data in db["bjts"].items():
    row = table.add_row().cells
    row[0].text = dev
    row[1].text = f"{data['w_um']:.4f}"
    row[2].text = f"{data['h_um']:.4f}"
    row[3].text = f"{data['area_um2']:.4f}"


doc.save('Comprehensive_Model_Report.docx')
print("Successfully generated Comprehensive_Model_Report.docx")
